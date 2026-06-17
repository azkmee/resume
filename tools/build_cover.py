#!/usr/bin/env python3
"""Build a clean cover-letter DOCX + one-page PDF to match the resume.

Single column, standard font, no graphics. Contact via env vars so personal
details never get committed to this public repo:
    RESUME_NAME, RESUME_EMAIL, RESUME_PHONE, RESUME_LINKEDIN, RESUME_LOCATION

Usage:
    RESUME_EMAIL=you@example.com python3 tools/build_cover.py
"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor

OUT_DIR = "build"
OUT_NAME = "Azmi-Cover-Letter-Mastercard"

NAME = os.environ.get("RESUME_NAME", "Azmi")
LOCATION = os.environ.get("RESUME_LOCATION", "Singapore")
contact_bits = [LOCATION]
for var in ("RESUME_EMAIL", "RESUME_PHONE", "RESUME_LINKEDIN"):
    v = os.environ.get(var)
    if v:
        contact_bits.append(v)
CONTACT = "  |  ".join(contact_bits)
DATED = date.today().strftime("%d %B %Y")

SALUTATION = "Dear Foundry R&D Hiring Team,"
SUBJECT = "Re: Senior Software Engineer (Generative AI) — Foundry R&D, Singapore"

PARAS = [
    "I'm applying for the Senior Software Engineer (Generative AI) role on the Mastercard "
    "Foundry R&D team. I'm a Singapore-based backend engineer with around five years at "
    "ESGpedia, where I build scalable Java/Spring Boot and Python services for an ESG data "
    "and certification platform — and the shape of this role, backend depth plus AI "
    "integration in an experiment-driven R&D setting, is exactly where I want to be.",

    "The part of the job description about integrating models behind reliable service "
    "interfaces — building service interfaces, managing data formats, and connecting "
    "external APIs — describes work I've already done. I integrated a managed machine "
    "learning service (AWS Textract) into our backend by designing a custom adapter, so "
    "model outputs slotted cleanly into downstream processing without coupling the rest of "
    "the system to the vendor's response shape. That same clean boundary between a model "
    "service and the application around it is what I'd bring to productionizing "
    "generative-AI features. Alongside it, I've shipped the backend fundamentals the role "
    "asks for: RESTful APIs with per-group access scoping, asynchronous processing and "
    "hot-reload caching on a data-intensive path (cutting around three seconds per update), "
    "and Flyway database migrations across four SQL and NoSQL databases auto-running in our "
    "deployment pipeline.",

    "On the generative-AI side I'll be straight: my production experience is in integrating "
    "machine learning services rather than training or prompt-engineering large language "
    "models at scale. What I do have is genuine momentum — I'm building a generative-AI "
    "workflow that generates and maintains our in-repo engineering documentation, and I "
    "learn fastest by building. The curious, adaptable, experiment-driven framing of this "
    "role is the environment I do my best work in.",

    "Two things may also be relevant: our platform is built in partnership with the "
    "Monetary Authority of Singapore, so I'm comfortable in regulated, finance-adjacent "
    "contexts; and I now lead and mentor a small remote backend team, running code reviews "
    "and pair-designed technical specifications — so the mentoring and best-practices part "
    "of the role is already how I work.",

    "I'd welcome the chance to talk. Thank you for your consideration.",
]

FONT = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)


def build_docx():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.7)
        s.left_margin = s.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.12

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(NAME); r.bold = True; r.font.size = Pt(18); r.font.name = FONT
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
    p.add_run(CONTACT).font.size = Pt(9.5)

    doc.add_paragraph(DATED).paragraph_format.space_after = Pt(8)
    sp = doc.add_paragraph(); sp.add_run(SUBJECT).bold = True
    doc.add_paragraph(SALUTATION)
    for para in PARAS:
        doc.add_paragraph(para)
    doc.add_paragraph("Sincerely,").paragraph_format.space_after = Pt(2)
    doc.add_paragraph(NAME)

    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, OUT_NAME + ".docx")
    doc.save(path)
    print("Wrote", path)
    return path


def build_pdf():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    def esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, OUT_NAME + ".pdf")
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        topMargin=1.6 * cm, bottomMargin=1.2 * cm, leftMargin=1.8 * cm, rightMargin=1.8 * cm,
        title=f"{NAME} — Cover Letter", author=NAME,
    )
    ss = getSampleStyleSheet()
    name_s = ParagraphStyle("n", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=18, leading=20, spaceAfter=1)
    contact_s = ParagraphStyle("c", parent=ss["Normal"], fontName="Helvetica", fontSize=9.5, leading=12, spaceAfter=10)
    body_s = ParagraphStyle("b", parent=ss["Normal"], fontName="Helvetica", fontSize=10.5, leading=14, spaceAfter=8, alignment=TA_JUSTIFY)
    subj_s = ParagraphStyle("s", parent=body_s, fontName="Helvetica-Bold", spaceAfter=8)

    story = [
        Paragraph(esc(NAME), name_s),
        Paragraph(esc(CONTACT), contact_s),
        Paragraph(esc(DATED), body_s),
        Paragraph(esc(SUBJECT), subj_s),
        Paragraph(esc(SALUTATION), body_s),
    ]
    for para in PARAS:
        story.append(Paragraph(esc(para), body_s))
    story += [Spacer(1, 6), Paragraph("Sincerely,", body_s), Paragraph(esc(NAME), body_s)]
    doc.build(story)
    print("Wrote", path, "| pages:", doc.page)
    if doc.page != 1:
        print("  WARNING: cover letter is not one page — trim content.")
    return path


if __name__ == "__main__":
    build_docx()
    build_pdf()
