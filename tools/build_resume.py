#!/usr/bin/env python3
"""Build an ATS-friendly, single-column resume DOCX (then convert to PDF via LibreOffice).

Design choices follow docs/RESUME-ATS-GUIDE.md: one column, standard headings, standard
font, consistent dates, no tables/graphics. Contact details come from environment
variables so personal info is never committed to this public repo:

    RESUME_NAME, RESUME_EMAIL, RESUME_PHONE, RESUME_LINKEDIN, RESUME_GITHUB, RESUME_LOCATION

Usage:
    RESUME_EMAIL=you@example.com python3 tools/build_resume.py
    soffice --headless --convert-to pdf --outdir build build/<file>.docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = "build"
OUT_NAME = "Azmi-Senior-Software-Engineer-Mastercard"

# ---- Contact (from env; placeholders otherwise) -----------------------------
NAME = os.environ.get("RESUME_NAME", "Azmi")
LOCATION = os.environ.get("RESUME_LOCATION", "Singapore")
contact_bits = [LOCATION]
for var in ("RESUME_EMAIL", "RESUME_PHONE", "RESUME_LINKEDIN", "RESUME_GITHUB"):
    val = os.environ.get(var)
    if val:
        contact_bits.append(val)
CONTACT = "  |  ".join(contact_bits)
HEADLINE = "Senior Software Engineer — Backend & Full-Stack"

SUMMARY = (
    "Software engineer with ~5 years building scalable backend services in Java (Spring "
    "Boot) and Python for a regulated ESG data and certification platform delivered in "
    "partnership with the Monetary Authority of Singapore. Strong in RESTful Application "
    "Programming Interface (API) design, microservices, asynchronous processing, caching, "
    "and SQL/NoSQL data-intensive systems on Amazon Web Services (AWS). Experienced "
    "integrating Machine Learning (ML) services into production backends, and actively "
    "building with generative Artificial Intelligence (AI) tooling. Now lead and mentor a "
    "remote backend team."
)

SKILLS = [
    ("Languages", "Java, Python, JavaScript (Node.js)"),
    ("Backend", "Spring Boot, RESTful API design, microservices, asynchronous and concurrent processing"),
    ("Data", "MySQL, MongoDB, Redis (caching), relational and NoSQL schema design, Flyway database migrations"),
    ("Cloud & DevOps", "Amazon Web Services (AWS) — Textract, RDS, S3; Continuous Integration/Continuous Delivery (CI/CD)"),
    ("AI / ML", "Machine Learning service integration (AWS Textract); generative AI tooling"),
    ("Testing", "JUnit, unit and integration testing"),
    ("Ways of working", "Agile/Scrum, sprint planning, code review, mentoring, distributed teams"),
]

EXPERIENCE = [
    ("Software Engineering Lead — ESGpedia, Singapore", "Apr 2026 – Present", [
        "Lead and mentor a remote backend team across Singapore, Indonesia, and the Philippines, owning sprint planning, story-point estimation, and developer briefings with the product manager; sustained zero attrition since taking over the team.",
        "Drive engineering quality through code reviews and pair-designed Technical Design Documents, coaching the senior engineer to autonomous ownership of multiple modules.",
        "Built a generative AI workflow that generates and maintains in-repository engineering documentation, reducing documentation-maintenance effort.",
    ]),
    ("Senior Software Engineer — ESGpedia, Singapore", "Apr 2025 – Apr 2026", [
        "Architected and delivered an ESG question-customisation module end to end — data model, RESTful APIs, persistence, and per-asset and per-user-group access scoping — so each user is served only the data relevant to their scope.",
        "Built a full-stack Apache Superset to AWS S3 integration using a backend endpoint that resolves document identifiers to short-lived presigned URLs, exposing documents to business-intelligence users without exposing storage paths.",
    ]),
    ("Software Engineer — ESGpedia, Singapore", "Jul 2022 – May 2025", [
        "Built an on-the-fly emissions calculator with asynchronous processing and a hot-reload caching layer on a data-intensive path, cutting ~3 seconds per update while keeping cached values fresh as upstream data changed.",
        "Integrated a managed ML service (AWS Textract optical character recognition) into backend processing behind a custom adapter, shipped as a one-week minimum viable product in Python, decoupling the system from the vendor's response format.",
        "Rolled out Flyway database migrations across 4 connections (3 SQL and 1 MongoDB), auto-running on every deployment to eliminate manual patching and produce a complete migration audit trail.",
    ]),
]

EDUCATION = [
    "B.E. (Honours), Engineering Systems and Design — Singapore University of Technology and Design (SUTD), 2017–2020",
    "Software and Product Development Certificate — SUTD, 2021",
]

FONT = "Calibri"
INK = RGBColor(0x1A, 0x1A, 0x1A)


def build():
    doc = Document()
    # tight margins so it stays one page
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    def heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = FONT

    # Name
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(NAME)
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = FONT
    # Headline + contact
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.add_run(HEADLINE).font.size = Pt(10.5)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.add_run(CONTACT).font.size = Pt(9.5)

    heading("Summary")
    doc.add_paragraph(SUMMARY)

    heading("Technical Skills")
    for label, items in SKILLS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(items)

    heading("Professional Experience")
    for title, dates, bullets in EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        p.add_run(title).bold = True
        d = doc.add_paragraph()
        d.paragraph_format.space_after = Pt(1)
        di = d.add_run(dates)
        di.italic = True
        di.font.size = Pt(9.5)
        for b in bullets:
            bp = doc.add_paragraph(b, style="List Bullet")
            bp.paragraph_format.space_after = Pt(2)

    heading("Education")
    for e in EDUCATION:
        doc.add_paragraph(e, style="List Bullet").paragraph_format.space_after = Pt(2)

    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, OUT_NAME + ".docx")
    doc.save(path)
    print("Wrote", path)
    return path


def _esc(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf():
    """Text-based, single-column A4 PDF via reportlab (parser-readable)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
    )

    os.makedirs(OUT_DIR, exist_ok=True)
    path = os.path.join(OUT_DIR, OUT_NAME + ".pdf")
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        topMargin=1.2 * cm, bottomMargin=1.0 * cm,
        leftMargin=1.4 * cm, rightMargin=1.4 * cm,
        title=f"{NAME} — Resume", author=NAME,
    )
    ss = getSampleStyleSheet()
    name_s = ParagraphStyle("name", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=20, leading=22, spaceAfter=1)
    headline_s = ParagraphStyle("headline", parent=ss["Normal"], fontName="Helvetica", fontSize=10.5, leading=13, spaceAfter=1)
    contact_s = ParagraphStyle("contact", parent=ss["Normal"], fontName="Helvetica", fontSize=9.5, leading=12, spaceAfter=6)
    head_s = ParagraphStyle("head", parent=ss["Normal"], fontName="Helvetica-Bold", fontSize=11, leading=13, spaceBefore=8, spaceAfter=3)
    body_s = ParagraphStyle("body", parent=ss["Normal"], fontName="Helvetica", fontSize=10, leading=12.5, spaceAfter=3, alignment=TA_LEFT)
    bullet_s = ParagraphStyle("bullet", parent=body_s, spaceAfter=2)
    role_s = ParagraphStyle("role", parent=body_s, fontName="Helvetica-Bold", spaceBefore=4, spaceAfter=0)
    dates_s = ParagraphStyle("dates", parent=ss["Normal"], fontName="Helvetica-Oblique", fontSize=9.5, leading=11, spaceAfter=2)

    def bullets(items):
        return ListFlowable(
            [ListItem(Paragraph(_esc(b), bullet_s), leftIndent=10, value="•") for b in items],
            bulletType="bullet", start="•", leftIndent=10, bulletFontSize=9,
        )

    story = [
        Paragraph(_esc(NAME), name_s),
        Paragraph(_esc(HEADLINE), headline_s),
        Paragraph(_esc(CONTACT), contact_s),
        Paragraph("SUMMARY", head_s), Paragraph(_esc(SUMMARY), body_s),
        Paragraph("TECHNICAL SKILLS", head_s),
    ]
    for label, items in SKILLS:
        story.append(Paragraph(f"<b>{_esc(label)}:</b> {_esc(items)}", body_s))
    story.append(Paragraph("PROFESSIONAL EXPERIENCE", head_s))
    for title, dates, items in EXPERIENCE:
        story.append(Paragraph(_esc(title), role_s))
        story.append(Paragraph(_esc(dates), dates_s))
        story.append(bullets(items))
    story.append(Paragraph("EDUCATION", head_s))
    story.append(bullets(EDUCATION))

    doc.build(story)
    print("Wrote", path, "| pages:", doc.page)
    if doc.page != 1:
        print("  WARNING: resume is not exactly one page — trim content.")
    return path


if __name__ == "__main__":
    build()
    build_pdf()
